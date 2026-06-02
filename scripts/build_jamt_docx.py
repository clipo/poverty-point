#!/usr/bin/env python3
"""
Build the JAMT Poverty Point Word documents from Markdown source.

Builds docs/jamt/Manuscript.md -> Manuscript.docx and
docs/jamt/Supplemental.md -> Supplemental.docx (figures already numbered;
no figure-swap or section-renumber preprocessing needed).

Steps per file:
1. Pre-process the markdown (absolute figure paths, em-dash cleanup, strip rules)
2. Convert via pandoc with the reference template (equation rendering)
3. Post-process with python-docx (fonts, heading styles, caption styling, image sizing)
"""

import re
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn

PROJECT_ROOT = Path('/Users/clipo/PycharmProjects/poverty-point')
JAMT_DIR = PROJECT_ROOT / 'docs' / 'jamt'
FIGURES_DIR = PROJECT_ROOT / 'figures'
REFERENCE_TEMPLATE = JAMT_DIR / 'reference_template.docx'
TEMP_MD = JAMT_DIR / '_temp_build.md'

TARGETS = [
    (JAMT_DIR / 'Manuscript.md', JAMT_DIR / 'Manuscript.docx'),
    (JAMT_DIR / 'Supplemental.md', JAMT_DIR / 'Supplemental.docx'),
]


def preprocess_markdown(md_source):
    """Pre-process the markdown: absolute figure paths, em-dash cleanup, strip rules."""
    content = Path(md_source).read_text(encoding='utf-8')

    # Convert relative figure paths to absolute for pandoc. Replace the longer
    # prefix first so the shorter one does not corrupt it.
    content = content.replace('../../figures/', str(FIGURES_DIR) + '/')
    content = content.replace('../figures/', str(FIGURES_DIR) + '/')

    # Safety net: convert any stray em-dashes (house style: none) to commas.
    content = content.replace('—', ', ')
    # En-dashes (U+2013) only when not a numeric range separator.
    content = re.sub(r'(?<!\d)–(?!\d)', ', ', content)

    # Remove thematic-break horizontal rules (ugly separators in Word).
    content = re.sub(r'\n---\n', '\n\n', content)

    TEMP_MD.write_text(content, encoding='utf-8')
    print("  markdown pre-processed.")


def convert_with_pandoc(output_docx):
    """Convert the temp markdown to docx using pandoc with the reference template."""
    cmd = [
        'pandoc', str(TEMP_MD),
        '-o', str(output_docx),
        '--reference-doc', str(REFERENCE_TEMPLATE),
        '--resource-path', str(JAMT_DIR),
        '--from', 'markdown+tex_math_dollars+raw_tex',
        '--to', 'docx',
        '--wrap=none',
    ]
    result = subprocess.run(cmd, capture_output=True, text=True)
    if result.returncode != 0:
        print(f"Pandoc error: {result.stderr}")
        sys.exit(1)
    TEMP_MD.unlink(missing_ok=True)
    print("  pandoc conversion completed.")


def post_process_docx(output_docx):
    """Post-process the generated docx for proper formatting."""
    doc = Document(str(output_docx))

    # 1. Times New Roman 11pt black for all body text.
    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.name is None:
                run.font.name = 'Times New Roman'
            if run.font.size is None:
                run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0, 0, 0)

    # 2. Heading styles per CLAUDE.md: H1 bold, H2 italic, H3 underline, H4 italic.
    heading_fmt = {
        'Heading 1': dict(size=14, bold=True, italic=False, underline=False),
        'Heading 2': dict(size=12, bold=False, italic=True, underline=False),
        'Heading 3': dict(size=11, bold=False, italic=False, underline=True),
        'Heading 4': dict(size=11, bold=False, italic=True, underline=False),
    }
    for para in doc.paragraphs:
        fmt = heading_fmt.get(para.style.name)
        if fmt:
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(fmt['size'])
                run.font.bold = fmt['bold']
                run.font.italic = fmt['italic']
                run.font.underline = fmt['underline']
                run.font.color.rgb = RGBColor(0, 0, 0)

    # 3. Figure captions: bold+italic first sentence, italic remainder, 10pt.
    # Format in place only. Bold-first-sentence and italics already come from the
    # markdown ***Figure N. ...*** / *...* markup, so we set size/font and never
    # rebuild run text from paragraph.text (which omits inline OMML math and would
    # orphan/clump the equations and leave empty gaps where symbols belong).
    for para in doc.paragraphs:
        text = para.text.strip()
        if re.match(r'^Figure S?\d+\.', text) and len(text) > 30:
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0, 0, 0)

    # 4. Size images to fit page width (6.0 inches).
    max_width = Inches(6.0)
    for para in doc.paragraphs:
        for run in para.runs:
            for child in run._element:
                if child.tag.endswith('}drawing'):
                    for ext in list(child.iter(qn('wp:extent'))) + list(child.iter(qn('a:ext'))):
                        cx = int(ext.get('cx', 0))
                        cy = int(ext.get('cy', 0))
                        if cx > max_width:
                            ratio = max_width / cx
                            ext.set('cx', str(int(cx * ratio)))
                            ext.set('cy', str(int(cy * ratio)))

    # 5. Final em-dash cleanup in the generated docx.
    for para in doc.paragraphs:
        for run in para.runs:
            if run.text and '—' in run.text:
                run.text = run.text.replace('—', ', ')

    # 6. Style any pandoc "Image Caption" paragraphs.
    for para in doc.paragraphs:
        if para.style.name == 'Image Caption' and para.text.strip():
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
                run.font.italic = True
                run.font.color.rgb = RGBColor(0, 0, 0)

    doc.save(str(output_docx))
    print(f"  post-processed and saved: {output_docx.name}")


def verify_document(output_docx):
    doc = Document(str(output_docx))
    eq = sum(len(p._element.findall('.//' + qn('m:oMath'))) for p in doc.paragraphs)
    img = sum(1 for rel in doc.part.rels.values() if 'image' in rel.reltype)
    em = sum(1 for p in doc.paragraphs if '—' in p.text)
    print(f"  VERIFY {output_docx.name}: equations={eq}, images={img}, em-dash paragraphs={em}, paragraphs={len(doc.paragraphs)}")


def build(md_source, output_docx):
    print(f"Building {output_docx.name} from {md_source.name}...")
    preprocess_markdown(md_source)
    convert_with_pandoc(output_docx)
    post_process_docx(output_docx)
    verify_document(output_docx)


if __name__ == '__main__':
    print("=" * 60)
    print("Building JAMT Poverty Point Word documents")
    print("=" * 60)
    for md_source, output_docx in TARGETS:
        build(md_source, output_docx)
    print("\nBuild complete.")
