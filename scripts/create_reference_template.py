#!/usr/bin/env python3
"""Create a Word reference template with proper styles for pandoc conversion."""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# Set default font for Normal style
normal = doc.styles['Normal']
normal.font.name = 'Times New Roman'
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor(0, 0, 0)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15

# Body Text style
body = doc.styles['Body Text']
body.font.name = 'Times New Roman'
body.font.size = Pt(11)
body.font.color.rgb = RGBColor(0, 0, 0)
body.paragraph_format.space_after = Pt(6)
body.paragraph_format.line_spacing = 1.15

# First Paragraph style (pandoc uses this for the first paragraph after headings)
from docx.enum.style import WD_STYLE_TYPE
try:
    first_para = doc.styles['First Paragraph']
except KeyError:
    first_para = doc.styles.add_style('First Paragraph', WD_STYLE_TYPE.PARAGRAPH)
first_para.font.name = 'Times New Roman'
first_para.font.size = Pt(11)
first_para.font.color.rgb = RGBColor(0, 0, 0)
first_para.paragraph_format.space_after = Pt(6)
first_para.paragraph_format.line_spacing = 1.15

# Heading 1: Bold
h1 = doc.styles['Heading 1']
h1.font.name = 'Times New Roman'
h1.font.size = Pt(14)
h1.font.bold = True
h1.font.italic = False
h1.font.underline = False
h1.font.color.rgb = RGBColor(0, 0, 0)
h1.paragraph_format.space_before = Pt(18)
h1.paragraph_format.space_after = Pt(6)

# Heading 2: Italic
h2 = doc.styles['Heading 2']
h2.font.name = 'Times New Roman'
h2.font.size = Pt(12)
h2.font.bold = False
h2.font.italic = True
h2.font.underline = False
h2.font.color.rgb = RGBColor(0, 0, 0)
h2.paragraph_format.space_before = Pt(14)
h2.paragraph_format.space_after = Pt(6)

# Heading 3: Underlined
h3 = doc.styles['Heading 3']
h3.font.name = 'Times New Roman'
h3.font.size = Pt(11)
h3.font.bold = False
h3.font.italic = False
h3.font.underline = True
h3.font.color.rgb = RGBColor(0, 0, 0)
h3.paragraph_format.space_before = Pt(12)
h3.paragraph_format.space_after = Pt(6)

# Heading 4
h4 = doc.styles['Heading 4']
h4.font.name = 'Times New Roman'
h4.font.size = Pt(11)
h4.font.bold = False
h4.font.italic = True
h4.font.underline = False
h4.font.color.rgb = RGBColor(0, 0, 0)

# Caption style
caption = doc.styles['Caption']
caption.font.name = 'Times New Roman'
caption.font.size = Pt(10)
caption.font.italic = True
caption.font.color.rgb = RGBColor(0, 0, 0)

# Image Caption style (custom pandoc style)
try:
    img_caption = doc.styles['Image Caption']
except KeyError:
    img_caption = doc.styles.add_style('Image Caption', WD_STYLE_TYPE.PARAGRAPH)
img_caption.font.name = 'Times New Roman'
img_caption.font.size = Pt(10)
img_caption.font.italic = True
img_caption.font.color.rgb = RGBColor(0, 0, 0)

# Title style
title = doc.styles['Title']
title.font.name = 'Times New Roman'
title.font.size = Pt(16)
title.font.bold = True
title.font.color.rgb = RGBColor(0, 0, 0)

# Subtitle
subtitle = doc.styles['Subtitle']
subtitle.font.name = 'Times New Roman'
subtitle.font.size = Pt(12)
subtitle.font.italic = True
subtitle.font.color.rgb = RGBColor(0, 0, 0)

# Abstract style
try:
    abstract = doc.styles['Abstract']
except KeyError:
    abstract = doc.styles.add_style('Abstract', WD_STYLE_TYPE.PARAGRAPH)
abstract.font.name = 'Times New Roman'
abstract.font.size = Pt(10)
abstract.font.italic = True
abstract.font.color.rgb = RGBColor(0, 0, 0)

# Set page margins
for section in doc.sections:
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

# Add Author style
try:
    author_style = doc.styles['Author']
except KeyError:
    author_style = doc.styles.add_style('Author', WD_STYLE_TYPE.PARAGRAPH)
author_style.font.name = 'Times New Roman'
author_style.font.size = Pt(11)
author_style.font.color.rgb = RGBColor(0, 0, 0)

# Add placeholder content so styles are used
doc.add_paragraph('Title', 'Title')
doc.add_paragraph('Subtitle', 'Subtitle')
doc.add_paragraph('Author', 'Author')
doc.add_paragraph('Abstract', 'Abstract')
doc.add_heading('Heading 1', level=1)
doc.add_heading('Heading 2', level=2)
doc.add_heading('Heading 3', level=3)
doc.add_heading('Heading 4', level=4)
doc.add_paragraph('Body text paragraph.', 'Body Text')
doc.add_paragraph('First paragraph after heading.', 'First Paragraph')
doc.add_paragraph('Caption text.', 'Caption')
doc.add_paragraph('Image Caption text.', 'Image Caption')

output_path = '/Users/clipo/PycharmProjects/poverty-point/docs/manuscript/reference_template.docx'
doc.save(output_path)
print(f"Reference template saved to {output_path}")
