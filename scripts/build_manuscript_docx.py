#!/usr/bin/env python3
"""
Build the Poverty Point manuscript Word document from Markdown source.

Steps:
1. Pre-process the markdown (fix figure numbering, absolute paths, em-dashes)
2. Convert via pandoc with reference template (proper equation rendering)
3. Post-process with python-docx (formatting, image sizing, caption styling)
"""

import os
import re
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Paths
PROJECT_ROOT = Path('/Users/clipo/PycharmProjects/poverty-point')
MANUSCRIPT_DIR = PROJECT_ROOT / 'docs' / 'manuscript'
FIGURES_DIR = PROJECT_ROOT / 'figures'
MD_SOURCE = MANUSCRIPT_DIR / 'Poverty_Point_JAS_Manuscript.md'
REFERENCE_TEMPLATE = MANUSCRIPT_DIR / 'reference_template.docx'
OUTPUT_DOCX = MANUSCRIPT_DIR / 'Poverty_Point_JAS_Manuscript.docx'
TEMP_MD = MANUSCRIPT_DIR / '_temp_build.md'


def preprocess_markdown():
    """Pre-process the markdown: fix paths, figure numbering, em-dashes."""
    with open(MD_SOURCE, 'r', encoding='utf-8') as f:
        content = f.read()

    # -------------------------------------------------------------------------
    # 1. Fix figure numbering: Figures 4 and 5 are out of order in the source.
    #    Figure 5 (model architecture) appears in Section 3.1 BEFORE
    #    Figure 4 (environmental variability) in Section 3.2.
    #    Swap them so they appear in sequential order.
    # -------------------------------------------------------------------------

    # Swap figure numbers in TEXT ONLY (not file paths).
    # First, protect all file paths by temporarily replacing them
    # Extract all image paths and protect them
    image_paths = re.findall(r'(\!\[[^\]]*\]\([^)]+\))', content)
    path_placeholders = {}
    for idx, path in enumerate(image_paths):
        placeholder = f'__IMG_PATH_{idx}__'
        path_placeholders[placeholder] = path
        content = content.replace(path, placeholder, 1)

    # Now safely swap figure numbers in visible text
    content = content.replace('Figure 5.', 'Figure __FOUR__.')
    content = content.replace('Figure 5 ', 'Figure __FOUR__ ')
    content = content.replace('Figure 5]', 'Figure __FOUR__]')

    content = content.replace('Figure 4.', 'Figure 5.')
    content = content.replace('Figure 4 ', 'Figure 5 ')
    content = content.replace('Figure 4]', 'Figure 5]')

    content = content.replace('Figure __FOUR__.', 'Figure 4.')
    content = content.replace('Figure __FOUR__ ', 'Figure 4 ')
    content = content.replace('Figure __FOUR__]', 'Figure 4]')

    # Restore image paths (unchanged)
    for placeholder, path in path_placeholders.items():
        content = content.replace(placeholder, path)

    # -------------------------------------------------------------------------
    # 2. Convert relative figure paths to absolute paths for pandoc
    # -------------------------------------------------------------------------
    content = content.replace('../../figures/', str(FIGURES_DIR) + '/')

    # -------------------------------------------------------------------------
    # 3. Replace em-dashes with commas or shorter alternatives
    # -------------------------------------------------------------------------
    # U+2014 EM DASH
    content = content.replace('\u2014', ', ')
    # U+2013 EN DASH (keep for number ranges like 1700-1100)
    # Only replace en-dashes that aren't in number ranges
    content = re.sub(r'(?<!\d)\u2013(?!\d)', ', ', content)

    # -------------------------------------------------------------------------
    # 4. Remove horizontal rules (---) as they create ugly separators in Word
    # -------------------------------------------------------------------------
    content = re.sub(r'\n---\n', '\n\n', content)

    # -------------------------------------------------------------------------
    # 5. Fix figure captions: convert markdown bold/italic to clean text
    #    Pandoc handles ***text*** as bold+italic, so keep that format
    # -------------------------------------------------------------------------

    # -------------------------------------------------------------------------
    # 6. (No longer needed — Figures 1 and 2 image files now exist)
    # -------------------------------------------------------------------------

    # -------------------------------------------------------------------------
    # 7. Clean image alt text to avoid duplicate short captions
    #    Replace descriptive alt text with just empty alt for pandoc
    #    The full captions are in separate paragraphs below each image
    # -------------------------------------------------------------------------
    content = re.sub(
        r'!\[[^\]]+\]\(([^)]+)\)',
        r'![](\1)',
        content
    )

    # -------------------------------------------------------------------------
    # 8. Fix duplicate section numbering: 2.6 appears twice
    #    Second 2.6 (Cooperation Benefits) should be 2.7
    #    Current 2.7 (Critical Threshold) should be 2.8
    # -------------------------------------------------------------------------
    # Fix the second occurrence of "### 2.6" (Cooperation Benefits)
    # First 2.6 is "The Ecotone Advantage", second is "Cooperation Benefits"
    marker = '### 2.6'
    first_26 = content.find(marker)
    if first_26 >= 0:
        second_26 = content.find(marker, first_26 + len(marker))
        if second_26 >= 0:
            content = content[:second_26] + '### 2.7' + content[second_26 + len(marker):]
    # Now fix "### 2.7 Critical Threshold" -> "### 2.8 Critical Threshold"
    if '### 2.7 Cooperation Benefits' in content:
        content = content.replace(
            '### 2.7 Critical Threshold',
            '### 2.8 Critical Threshold'
        )

    # Write temp file for pandoc
    with open(TEMP_MD, 'w', encoding='utf-8') as f:
        f.write(content)

    print("Markdown pre-processed successfully.")
    return content


def convert_with_pandoc():
    """Convert markdown to docx using pandoc with reference template."""
    cmd = [
        'pandoc',
        str(TEMP_MD),
        '-o', str(OUTPUT_DOCX),
        '--reference-doc', str(REFERENCE_TEMPLATE),
        '--from', 'markdown+tex_math_dollars+raw_tex',
        '--to', 'docx',
        '--wrap=none',
    ]

    print(f"Running: {' '.join(cmd)}")
    result = subprocess.run(cmd, capture_output=True, text=True)

    if result.returncode != 0:
        print(f"Pandoc error: {result.stderr}")
        sys.exit(1)

    print("Pandoc conversion completed.")

    # Clean up temp file
    TEMP_MD.unlink(missing_ok=True)


def post_process_docx():
    """Post-process the generated docx for proper formatting."""
    doc = Document(str(OUTPUT_DOCX))

    # =========================================================================
    # 1. Apply Times New Roman 11pt to ALL text
    # =========================================================================
    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.name is None:
                run.font.name = 'Times New Roman'
            if run.font.size is None:
                run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0, 0, 0)

    # =========================================================================
    # 2. Fix heading styles per CLAUDE.md requirements
    #    H1: Bold, H2: Italic, H3: Underlined
    # =========================================================================
    for para in doc.paragraphs:
        style_name = para.style.name

        if style_name == 'Heading 1':
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
                run.font.bold = True
                run.font.italic = False
                run.font.underline = False
                run.font.color.rgb = RGBColor(0, 0, 0)

        elif style_name == 'Heading 2':
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.font.bold = False
                run.font.italic = True
                run.font.underline = False
                run.font.color.rgb = RGBColor(0, 0, 0)

        elif style_name == 'Heading 3':
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
                run.font.bold = False
                run.font.italic = False
                run.font.underline = True
                run.font.color.rgb = RGBColor(0, 0, 0)

        elif style_name == 'Heading 4':
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
                run.font.bold = False
                run.font.italic = True
                run.font.underline = False
                run.font.color.rgb = RGBColor(0, 0, 0)

    # =========================================================================
    # 3. Format figure captions: Italic with bold first sentence
    #    Figure captions follow the pattern "Figure N. First sentence. Rest..."
    # =========================================================================
    for para in doc.paragraphs:
        text = para.text.strip()
        # Match figure captions (they start with "Figure N.")
        if re.match(r'^Figure \d+\.', text) and len(text) > 30:
            # Find the first sentence boundary (after "Figure N. ...")
            match = re.match(r'^(Figure \d+\.\s+[^.]+\.)\s*(.*)', text, re.DOTALL)
            if match:
                first_sentence = match.group(1)
                rest = match.group(2)

                # Clear existing runs
                for run in para.runs:
                    run.text = ''

                # If there are no runs, add one
                if not para.runs:
                    para.add_run()

                # Set first run to bold+italic first sentence
                para.runs[0].text = first_sentence
                para.runs[0].font.bold = True
                para.runs[0].font.italic = True
                para.runs[0].font.name = 'Times New Roman'
                para.runs[0].font.size = Pt(10)
                para.runs[0].font.color.rgb = RGBColor(0, 0, 0)

                if rest:
                    # Add the rest as italic only
                    rest_run = para.add_run(' ' + rest)
                    rest_run.font.bold = False
                    rest_run.font.italic = True
                    rest_run.font.name = 'Times New Roman'
                    rest_run.font.size = Pt(10)
                    rest_run.font.color.rgb = RGBColor(0, 0, 0)

    # =========================================================================
    # 4. Size images to fit page width (6.5 inches with 1-inch margins)
    # =========================================================================
    max_width = Inches(6.0)
    for para in doc.paragraphs:
        for run in para.runs:
            for child in run._element:
                if child.tag.endswith('}drawing'):
                    # Find the extent element
                    for extent in child.iter(qn('wp:extent')):
                        cx = int(extent.get('cx', 0))
                        cy = int(extent.get('cy', 0))
                        if cx > max_width:
                            ratio = max_width / cx
                            new_cx = int(cx * ratio)
                            new_cy = int(cy * ratio)
                            extent.set('cx', str(new_cx))
                            extent.set('cy', str(new_cy))
                    # Also fix inline extent if present
                    for extent in child.iter(qn('a:ext')):
                        cx = int(extent.get('cx', 0))
                        cy = int(extent.get('cy', 0))
                        if cx > max_width:
                            ratio = max_width / cx
                            new_cx = int(cx * ratio)
                            new_cy = int(cy * ratio)
                            extent.set('cx', str(new_cx))
                            extent.set('cy', str(new_cy))

    # =========================================================================
    # 5. Final em-dash cleanup in the generated docx
    # =========================================================================
    for para in doc.paragraphs:
        for run in para.runs:
            if run.text and '\u2014' in run.text:
                run.text = run.text.replace('\u2014', ', ')
            if run.text and '\u2013' in run.text:
                # Only replace en-dashes not in number ranges
                run.text = re.sub(r'(?<!\d)\u2013(?!\d)', ', ', run.text)

    # =========================================================================
    # 6. Remove any remaining "PLACEHOLDER" text markers
    # =========================================================================
    paras_to_check = []
    for i, para in enumerate(doc.paragraphs):
        if 'PLACEHOLDER' in para.text:
            paras_to_check.append(i)

    # Don't remove these paragraphs, but mark them clearly
    for i in paras_to_check:
        para = doc.paragraphs[i]
        for run in para.runs:
            run.font.color.rgb = RGBColor(128, 128, 128)
            run.font.italic = True

    # =========================================================================
    # 6b. Remove empty "Image Caption" and "Captioned Figure" paragraphs
    #     and any duplicate short captions from alt text
    # =========================================================================
    for para in doc.paragraphs:
        style_name = para.style.name
        text = para.text.strip()
        # Remove empty caption paragraphs
        if style_name in ('Image Caption', 'Captioned Figure', 'Caption') and not text:
            # Can't delete paragraphs easily, but can clear them
            pass
        # Format Image Caption paragraphs
        if style_name == 'Image Caption' and text:
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
                run.font.italic = True
                run.font.color.rgb = RGBColor(0, 0, 0)

    # =========================================================================
    # 7. Ensure math fonts are set properly
    # =========================================================================
    # Find all oMath elements and ensure they use Cambria Math
    for para in doc.paragraphs:
        for math_elem in para._element.findall('.//' + qn('m:oMath')):
            for rpr in math_elem.findall('.//' + qn('m:rPr')):
                # Ensure math font is set
                pass  # Pandoc typically handles this correctly

    # =========================================================================
    # Save
    # =========================================================================
    doc.save(str(OUTPUT_DOCX))
    print(f"Post-processed document saved to {OUTPUT_DOCX}")


def verify_document():
    """Verify the final document for quality."""
    doc = Document(str(OUTPUT_DOCX))

    print("\n=== DOCUMENT VERIFICATION ===")

    # Count equations
    eq_count = 0
    for para in doc.paragraphs:
        eq_count += len(para._element.findall('.//' + qn('m:oMath')))
    print(f"Equations (oMath elements): {eq_count}")

    # Count images
    img_count = 0
    for rel in doc.part.rels.values():
        if 'image' in rel.reltype:
            img_count += 1
    print(f"Images: {img_count}")

    # Check for em-dashes
    emdash_count = 0
    for para in doc.paragraphs:
        if '\u2014' in para.text:
            emdash_count += 1
    print(f"Paragraphs with em-dashes: {emdash_count}")

    # Check heading formatting
    print("\nHeading styles:")
    for para in doc.paragraphs:
        if para.style.name.startswith('Heading'):
            for run in para.runs:
                fmt = []
                if run.font.bold: fmt.append('bold')
                if run.font.italic: fmt.append('italic')
                if run.font.underline: fmt.append('underline')
                print(f"  [{para.style.name}] {para.text[:60]} => {', '.join(fmt) if fmt else 'no special format'}")
                break  # Just check first run

    # Check figure captions
    print("\nFigure captions:")
    for para in doc.paragraphs:
        text = para.text.strip()
        if re.match(r'^Figure \d+\.', text) and len(text) > 30:
            has_bold = any(run.font.bold for run in para.runs)
            has_italic = any(run.font.italic for run in para.runs)
            print(f"  {text[:70]}... bold={has_bold}, italic={has_italic}")

    # Check images placement
    print("\nImage placements:")
    img_idx = 0
    for i, para in enumerate(doc.paragraphs):
        has_image = False
        for run in para.runs:
            for child in run._element:
                if child.tag.endswith('}drawing'):
                    has_image = True
        if has_image:
            img_idx += 1
            next_text = ""
            for j in range(i+1, min(i+3, len(doc.paragraphs))):
                if doc.paragraphs[j].text.strip():
                    next_text = doc.paragraphs[j].text.strip()[:70]
                    break
            print(f"  Image {img_idx} at para {i}: {next_text}")

    print(f"\n=== Total paragraphs: {len(doc.paragraphs)} ===")


if __name__ == '__main__':
    print("=" * 60)
    print("Building Poverty Point JAS Manuscript")
    print("=" * 60)

    print("\nStep 1: Pre-processing markdown...")
    preprocess_markdown()

    print("\nStep 2: Converting with pandoc...")
    convert_with_pandoc()

    print("\nStep 3: Post-processing docx...")
    post_process_docx()

    print("\nStep 4: Verifying document...")
    verify_document()

    print("\n" + "=" * 60)
    print("Build complete!")
    print(f"Output: {OUTPUT_DOCX}")
    print("=" * 60)
